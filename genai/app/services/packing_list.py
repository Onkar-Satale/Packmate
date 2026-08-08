# Service module managing LLM packing list generation via Groq and Word document creation

import json
import re
from io import BytesIO
from datetime import datetime
from fastapi import HTTPException
from groq import Groq
from docx import Document
from app.config.settings import GROQ_API_KEY, logger
from app.services.weather import compute_full_trip_weather


# Simple in-memory cache for generated packing lists to prevent AI spam attacks
generation_cache = {}

def is_section_heading(item: str) -> bool:
    item = item.strip()
    if not item:
        return False
    if item.isupper():
        return True
        
    return False

def create_docx(packing_list: list):
    """
    Creates an in-memory Microsoft Word document (.docx) containing the packing list.
    
    Args:
        packing_list (list): The list of packing items including category headers.
        
    Returns:
        BytesIO: A byte buffer containing the Word document data, ready for streaming.
    """
    buffer = BytesIO()
    doc = Document()
    
    # Add title and current date
    doc.add_heading("🎒 Smart Packing Assistant", level=1)
    
    current_date = datetime.now().strftime("%Y-%m-%d")
    doc.add_paragraph(f"Generated on: {current_date}")
            
    # Add packing list items
    for item in packing_list:
        item_stripped = item.strip()
        if not item_stripped:
            continue
        
        if is_section_heading(item_stripped):
            doc.add_heading(item_stripped, level=2)
        else:
            doc.add_paragraph(item_stripped, style="List Bullet")
        
    doc.save(buffer)
    buffer.seek(0) # Reset buffer position to the start for reading
    return buffer

def generate_packing_data(data: dict):
    """
    The core logic integrating with the Groq Large Language Model.
    Sends environmental and traveler data to generate a complete packing list.
    
    Args:
        data (dict): The dictionary representation of the TripRequestGenerate model.
        
    Returns:
        dict: A parsed JSON response matching the required structure:
            { "packing_list": ["header", "item", ...] }
    """
    # Initialize the Groq client to call the LLM
    client = Groq(api_key=GROQ_API_KEY)

    # Gather dynamic day-wise temperature array from precomputed exact API logic
    temp_info = compute_full_trip_weather(data)
    
    # Define the system prompt guiding the AI's behavior, establishing rules, and restricting the output format
    system_prompt = """
========================
SMART PACKING ASSISTANT - DETAILED AI-GENERATED PACKING LIST
========================

You are a senior professional travel planner and packing consultant. Your task is to generate a complete, structured, and practical packing list for any traveler or group based entirely on the provided trip and traveler information. 
⚠️ IMPORTANT:
- Always include **all 12 mandatory sections** listed below in the exact order, even if minimal items are needed.
- Quantities, emojis, and items must be **dynamically decided** based on input data.
- Use **standardized emojis per section** consistently.
- Consider **traveler-specific details**: age, gender, medical notes, dietary restrictions, chronic conditions.
- Consider **trip details**: destination, activities, duration, accommodation, budget, luggage style.
- Include optional, backup, and emergency items for all travelers.
- EXTREMELY IMPORTANT: Take the provided day-by-day temperatures VERY seriously. Evaluate the exact high and low bounds and precisely adjust clothing (thermal wear, summer wear, thin layers) to perfectly match the temperature swings. Do NOT recommend winter items for hot days, or summer items for freezing days.
- Ensure **practicality**: only include items travelers can realistically carry.
- Output **only the packing list in a flat JSON list**, do not include greetings, explanations, or nested arrays.


========================
MANDATORY SECTIONS (ALL CAPS)
========================
1. DOCUMENTS
2. CLOTHING
3. FOOTWEAR
4. TOILETRIES & PERSONAL CARE
5. ELECTRONICS & GADGETS
6. MEDICAL & HEALTH
7. ACCESSORIES
8. FOOD & SNACKS
9. ACTIVITY-SPECIFIC ITEMS
10. MISCELLANEOUS
11. WEATHER-SPECIFIC ITEMS
12. SAFETY & EMERGENCY ITEMS

========================
GENERAL INSTRUCTIONS
========================
- Include **quantities appropriate** for number of days, travelers, and laundry availability.
- Include **all traveler-specific adjustments**: age, gender, medical notes, dietary restrictions.
- Include **weather-appropriate items** based on destination temperature and season.
- Include **electronics and accessories** according to traveler details (phones, laptops, cameras, chargers).
- Include **food, snacks, and hydration** as per traveler type (solo, kids, elders, family).
- Include **backup and emergency items**: extra socks, ch+argers, first-aid, medications.
- For multi-person trips, include items for **each traveler individually**.
- **Do not use placeholders** like "stuff" or "if needed".
- Always use **consistent emojis per section**.
- Output must be a **flat JSON list**, not nested arrays.
- Ensure all items are **practical, ready-to-go, and realistic**.
- **Every section MUST have at least one item.** If no specific activities are provided, add general exploration items (e.g. "Daypack 🎒", "Comfortable walking wear 👕") to ACTIVITY-SPECIFIC ITEMS so it is never empty.

========================
OUTPUT FORMAT
========================
Respond ONLY with **JSON object** in this exact format:

{
  "packing_list": [
    "DOCUMENTS",
    "Passport 📄 (1 per traveler)",
    "Visa 🛂 (if required)",
    "Flight tickets 🎟️",
    "Hotel booking confirmation 🏨",
    "Travel insurance card 🏥",
    ...
    "CLOTHING",
    "T-shirts 👕 (N)",
    "Jeans / Pants 👖 (N)",
    "Shorts 🩳 (N)",
    "Sleepwear 😴 (N)",
    ...
    "FOOTWEAR",
    "Walking shoes 👟",
    "Sandals 🩴",
    ...
    "TOILETRIES & PERSONAL CARE",
    "Toothbrush 🪥",
    "Toothpaste 🪥",
    "Soap / Body wash 🧼",
    ...
    "MEDICAL & HEALTH",
    "First aid kit 🩹",
    "Prescription medication 💊",
    "Fever medication 💊",
    ...
    "ACCESSORIES",
    "Sunglasses 🕶️",
    "Backpack 🎒",
    "Umbrella ☔️",
    ...
    "ELECTRONICS & GADGETS",
    "Smartphone & charger 📱",
    "Laptop & charger 💻",
    "Camera 📷",
    ...
    "FOOD & SNACKS",
    "Non-perishable snacks 🍫",
    "Energy bars 🍿",
    "Reusable water bottle 💧",
    ...
    "ACTIVITY-SPECIFIC ITEMS",
    "Gym clothes 🏋️",
    "Swimwear 🏊",
    "Trekking shoes 🥾",
    ...
    "MISCELLANEOUS",
    "Notebook 📝",
    "Guidebook 📘",
    ...
    "WEATHER-SPECIFIC ITEMS",
    "Thermal wear 🧥 (cold weather)",
    "Sunscreen 🌞 (hot weather)",
    "Raincoat ☔️ (rainy weather)",
    ...
    "SAFETY & EMERGENCY ITEMS",
    "Travel locks 🔒",
    "Power bank 🔋",
    "Flashlight 🔦",
    "Emergency contact info 📇"
  ]
}
- Do not use nested arrays. Use a **flat list**.
- Each **section header is ALL CAPS, max 30 characters, no emojis**.
- Each **item under a section must include emoji** and quantity if relevant.
- Generate **all items dynamically**. Do not use placeholders like "if needed" or "stuff".
"""

    # Populate the dynamic data for the current request
    user_prompt = f"""
Destination: {data['destination']}
Duration: {data['days']} days
Trip Type: {data['trip_type']}
Purpose: {data['purpose']}
Activities: {data['activities']}
Stay Type: {data['stay_type']}
Budget: {data['budget']}
Food Preference: {data['food']}
Luggage Style: {data['luggage']}
Travel Mode: {data['travel_type']}
Travelers details: {data['travelers']}
Medical Notes: {data.get('medicalNotes')}
Dietary Notes: {data.get('dietaryNotes')}
Laundry Available: {data.get('laundry')}
Shopping Planned: {data.get('shopping')}
Photography Gear Needed: {data.get('photographyGear')}
Work Laptop: {data.get('workLaptop')}


Include all packing items dynamically based on:
- Traveler-specific medical notes, age, gender
- Weather conditions and temperature
- Trip duration, activities, and accommodation
- Budget and luggage style
- Ensure all 12 mandatory sections are included
- Use consistent emojis and realistic quantities

Environmental Context:
{temp_info}. Use this to determine weather-appropriate items.

Generate the JSON packing list now.
"""

    # Make the call to the Groq API model
    # We specify response_format={"type": "json_object"} to strictly enforce a JSON response structure.
    res = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        temperature=0.4,
        max_completion_tokens=2000,
        response_format={"type": "json_object"}
    )
    
    try:
        content = res.choices[0].message.content
        
        try:
            result = json.loads(content)
        except json.JSONDecodeError as exc:
            logger.error(f"JSON decode error from Groq: {exc}\nRaw Content: {content[:200]}")
            raise HTTPException(status_code=500, detail="Invalid JSON response from AI")
        
        # Ensure packing_list items are clean from stray bullet points or numbering if the AI still added them
        clean_list = []
        for line in result.get("packing_list", []):
            if not isinstance(line, str):
                continue
            # Regex to remove bullets or numbering at the start of a response item
            clean_line = re.sub(r"^[•\-*\d.]+\s*", "", line).strip()
            if clean_line:
                clean_list.append(clean_line)
                
        logger.info(f"Successfully generated packing list containing {len(clean_list)} items")
        return {"packing_list": clean_list}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed processing AI response: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate packing list: {str(e)}")
