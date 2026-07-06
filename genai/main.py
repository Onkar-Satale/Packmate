"""
Smart Packing Assistant API - FastAPI Backend

This Python fastAPI service acts as the GenAI backend for the Smart Packing Assistant application. 
It receives trip details from the frontend, queries an LLM (Groq API) for a customized packing list 
in JSON format, and returns the list. It also provides functionality to save trips to a separate Node.js 
MongoDB backend and generate downloadable DOCX files.
"""

# ==========================================
# ### IMPORTS & CONFIGURATION ###
# ==========================================
import os
import json
import time
from typing import Optional
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, Header, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel, Field
from io import BytesIO
from fastapi import Request
from docx import Document
import re
import requests
from groq import Groq
from pathlib import Path
from datetime import datetime, timedelta
import random
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

# Resolve base directory to locate the .env file containing API keys
BASE_DIR = Path(__file__).resolve().parent
ENV_PATH = BASE_DIR / ".env"

# Load environment variables from the .env file into the os.environ dictionary
load_dotenv(dotenv_path=ENV_PATH, override=True)

# Pre-load RAG travel_chatbot in a background thread to prevent blocking Uvicorn's port binding.
# This ensures that the application starts up instantly and binds to the port, preventing Render deployment timeouts.
import threading
def pre_load_rag():
    try:
        logger.info("Starting background pre-loading of RAG model and ChromaDB...")
        from knowledge_base.rag import travel_chatbot, get_collection
        get_collection()
        logger.info("Successfully pre-loaded RAG model and ChromaDB cache in the background.")
    except Exception as e:
        logger.error(f"Failed to pre-load RAG model in background: {e}")

threading.Thread(target=pre_load_rag, daemon=True).start()

# Retrieve API keys from environment variables
GROQ_API_KEY = os.getenv("GROQ_API_KEY")

# Ensure the Groq API key is present before starting the app.
if not GROQ_API_KEY:
    raise Exception("GROQ_API_KEY not loaded!")

GENAI_API_SECRET = os.getenv("GENAI_API_SECRET", "")

def verify_api_key(x_api_key: str = Header(None)):
    if not GENAI_API_SECRET:
        logger.warning("GENAI_API_SECRET is not set. Service is unprotected!")
    elif x_api_key != GENAI_API_SECRET:
        raise HTTPException(status_code=401, detail="Unauthorized: Invalid API Key")

# Optimize requests globally with a session
http_session = requests.Session()

# ==========================================
# ### APP INITIALIZATION & MIDDLEWARE ###
# ==========================================

# Simple in-memory cache for generated packing lists to prevent AI spam attacks
generation_cache = {}
# Cache for 10-day weather forecasts by location to prevent redundant API calls
weather_cache = {}

# Initialize the FastAPI application
app = FastAPI(title="🎒Smart Packing Assistant API")

# Rate limiting is now handled securely by the Node.js API Gateway (express-rate-limit).

# Health-check root endpoint
@app.get("/")
def root():
    """
    Root endpoint to verify the API is running successfully.
    Returns: A simple JSON message indicating live status.
    """
    return {"message": "Smart Packing Assistant API is live ✅"}

# Allow frontend origins completely defined dynamically via .env
frontend_env = os.getenv("FRONTEND_URL", "")
FRONTEND_URL = [url.strip() for url in frontend_env.split(",") if url.strip()]

# Add CORS middleware to allow the React frontend to communicate with this backend.
app.add_middleware(
    CORSMiddleware, 
    allow_origins=FRONTEND_URL,
    allow_credentials=True,
    allow_methods=["*"],  # Allows all HTTP methods (GET, POST, etc.)
    allow_headers=["*"],  # Allows all headers
)

# ==========================================
# ### PYDANTIC DATA MODELS ###
# ==========================================
# These models define the expected shape of the incoming JSON requests.
# FastAPI validates incoming requests automatically based on these definitions.

class TripRequestGenerate(BaseModel):
    """
    Expected input format for generating a packing list or downloading it.
    """
    location: str = Field(..., max_length=150)
    days: int = Field(..., ge=1, le=120)
    trip_type: str = Field(..., max_length=50)
    purpose: str = Field(..., max_length=50)
    activities: str = Field(..., max_length=300)
    stay_type: str = Field(..., max_length=50)
    budget: str = Field(..., max_length=50)
    food: str = Field(..., max_length=100)
    luggage: str = Field(..., max_length=100)
    travel_type: str = Field(..., max_length=100)
    people: str = Field(..., max_length=1000)  # A flat string describing all travelers (e.g. "John, 25 years, Male")
    temperature: Optional[float] = None
    start_date: Optional[str] = Field(None, max_length=20)
    end_date: Optional[str] = Field(None, max_length=20)

class PrefetchWeatherRequest(BaseModel):
    """
    Payload for fetching weather and correcting city names
    """
    location: str = Field(..., max_length=150)


class DownloadRequest(BaseModel):
    """
    Expected input format for downloading a generated packing list.
    """
    packing_list: list

# ==========================================
# ### HELPER FUNCTIONS ###
# ==========================================

def is_section_heading(item: str) -> bool:
    item = item.strip()
    if not item:
        return False
    if item.isupper():
        return True
        
    return False

def create_docx(packing_list: list):
    """
    Creates an in-memory Microsoft Word document (.docx) containing the packing list.
    
    Args:
        packing_list (list): The list of packing items including category headers.
        
    Returns:
        BytesIO: A byte buffer containing the Word document data, ready for streaming.
    """
    buffer = BytesIO()
    doc = Document()
    
    # Add title and current date
    doc.add_heading("🎒 Smart Packing Assistant", level=1)
    
    current_date = datetime.now().strftime("%Y-%m-%d")
    doc.add_paragraph(f"Generated on: {current_date}")
            
    # Add packing list items
    for item in packing_list:
        item_stripped = item.strip()
        if not item_stripped:
            continue
        
        if is_section_heading(item_stripped):
            doc.add_heading(item_stripped, level=2)
        else:
            doc.add_paragraph(item_stripped, style="List Bullet")
        
    doc.save(buffer)
    buffer.seek(0) # Reset buffer position to the start for reading
    return buffer

def geocode_location(location: str):
    """
    Looks up latitude and longitude for a given city string using Open-Meteo's Geocoding API.
    """
    try:
        url = f"https://geocoding-api.open-meteo.com/v1/search?name={location}&count=1"
        res = http_session.get(url, timeout=5)
        if res.status_code == 200:
            data = res.json()
            if "results" in data and len(data["results"]) > 0:
                loc_data = data["results"][0]
                lat = loc_data.get("latitude")
                lon = loc_data.get("longitude")
                logger.info(f"Geocoding success for '{location}': lat={lat}, lon={lon}")
                return lat, lon
            else:
                logger.warning(f"Geocoding found no results for '{location}'")
                return None, None
        else:
            logger.error(f"Geocoding API returned status {res.status_code}")
            return None, None
    except Exception as e:
        logger.error(f"Geocoding error: {e}")
        return None, None

def prefetch_and_cache_weather(location: str):
    """
    Fetches the 16-day weather forecast using Open-Meteo API and caches it.
    Returns the current/average temperature for the provided destination for response compatibility.
    """
    loc_key = location.lower().strip()
    lat, lon = geocode_location(loc_key)
    if lat is None or lon is None:
        logger.warning(f"Skipping Open-Meteo forecast fetch due to geocoding failure for '{location}'")
        return None

    try:
        url = f"https://api.open-meteo.com/v1/forecast?latitude={lat}&longitude={lon}&daily=temperature_2m_max,temperature_2m_min&forecast_days=16&timezone=auto"
        res = http_session.get(url, timeout=5)
        
        if res.status_code == 200:
            data = res.json()
            api_temps = {}
            current_temp = None
            last_api_temp = None
            
            if "daily" in data:
                daily = data["daily"]
                time_list = daily.get("time", [])
                max_temps = daily.get("temperature_2m_max", [])
                min_temps = daily.get("temperature_2m_min", [])
                
                for i, dt_key in enumerate(time_list):
                    t_max = max_temps[i]
                    t_min = min_temps[i]
                    if t_max is not None and t_min is not None:
                        avg_temp = round((t_max + t_min) / 2.0, 1)
                        api_temps[dt_key] = avg_temp
                        last_api_temp = avg_temp
                        if i == 0:
                            current_temp = avg_temp
                            
            if current_temp is None and last_api_temp is not None:
                current_temp = last_api_temp
                
            weather_cache[loc_key] = {
                "api_temps": api_temps,
                "last_api_temp": last_api_temp,
                "fetched_at": time.time()
            }
            
            logger.info(f"Weather prefetched from Open-Meteo for '{location}' and cached successfully.")
            
            return current_temp
        else:
            logger.error(f"Open-Meteo API /forecast returned status {res.status_code}")
            return None
    except Exception as e:
        logger.error(f"Open-Meteo API error: {e}")
        return None

def compute_full_trip_weather(data: dict) -> str:
    location = data.get("location", "").lower().strip()
    start_date_str = data.get("start_date")
    end_date_str = data.get("end_date")
    trip_days = int(data.get("days", 1))
    
    fallback_temp = data.get('temperature')
    if fallback_temp is None:
        fallback_temp = random.randint(20, 30)
        logger.info(f"No temp provided, using random fallback: {fallback_temp}°C")

    try:
        # Fallback if dates are missing: use today + trip_days
        if not start_date_str or not end_date_str:
            logger.info(f"Dates missing for {location}, falling back to today + {trip_days} days")
            start_date = datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
            total_days = trip_days
        else:
            start_date = datetime.strptime(start_date_str, "%Y-%m-%d")
            end_date = datetime.strptime(end_date_str, "%Y-%m-%d")
            total_days = (end_date - start_date).days + 1
            if total_days <= 0: total_days = 1
        
        date_list = [start_date + timedelta(days=i) for i in range(total_days)]
        
        cached_weather = weather_cache.get(location)
        if not cached_weather:
            # Graceful fetch in case prefetch wasn't called or cache was lost
            logger.info("Weather cache miss in compute_full_trip_weather, fetching now.")
            _ = prefetch_and_cache_weather(location)
            cached_weather = weather_cache.get(location)
            
        if cached_weather and cached_weather.get("api_temps"):
            api_temps = cached_weather["api_temps"]
            # Retrieve 16th day (last available) temperature to act as our base
            sorted_dates = sorted(api_temps.keys())
            sixteenth_day_temp = api_temps[sorted_dates[-1]] if sorted_dates else fallback_temp
        else:
            api_temps = {}
            sixteenth_day_temp = fallback_temp

        forecast_api_lines = []
        forecast_extra_lines = []
        current_temp = sixteenth_day_temp

        for dt in date_list:
            dt_str = dt.strftime("%Y-%m-%d")
            
            if dt_str in api_temps:
                current_temp = api_temps[dt_str]
                forecast_api_lines.append(f"{dt_str} → {current_temp}°C (Prefetched)")
            else:
                drift = random.choices([-2, -1, 0, 1, 2], weights=[0.1, 0.35, 0.1, 0.35, 0.1])[0]
                new_temp = current_temp + drift

                if new_temp > sixteenth_day_temp + 5:
                    new_temp = sixteenth_day_temp + 5
                elif new_temp < sixteenth_day_temp - 5:
                    new_temp = sixteenth_day_temp - 5

                current_temp = round(new_temp, 1)
                forecast_extra_lines.append(f"{dt_str} → {current_temp}°C (Generated Drift)")
        
        all_lines = forecast_api_lines + forecast_extra_lines
        
        formatted_output = f"\n================ FULL DAY-WISE TEMPERATURE MAPPING FOR '{location.upper()}' ================\n"
        formatted_output += "\n".join(all_lines)
        formatted_output += "\n=================================================================================="
        
        logger.info(formatted_output)
        
        clean_lines = [line.replace(" (Prefetched)", "").replace(" (Generated Drift)", "") for line in all_lines]
        weather_text = "Day-wise temperature forecast:\n" + "\n".join(clean_lines)

        return weather_text

    except Exception as e:
        logger.error(f"Error generating full trip weather: {e}")
        return f"Average temperature: {fallback_temp}°C"

def generate_packing_data(data: dict):
    """
    The core logic integrating with the Groq Large Language Model.
    Sends environmental and traveler data to generate a complete packing list.
    
    Args:
        data (dict): The dictionary representation of the TripRequestGenerate model.
        
    Returns:
        dict: A parsed JSON response matching the required structure:
            { "packing_list": ["header", "item", ...] }
    """
    # Initialize the Groq client to call the LLM
    client = Groq(api_key=GROQ_API_KEY)

    # Gather dynamic day-wise temperature array from precomputed exact API logic
    temp_info = compute_full_trip_weather(data)
    
    # Define the system prompt guiding the AI's behavior, establishing rules, and restricting the output format
    system_prompt = """
========================
SMART PACKING ASSISTANT - DETAILED AI-GENERATED PACKING LIST
========================

You are a senior professional travel planner and packing consultant. Your task is to generate a complete, structured, and practical packing list for any traveler or group based entirely on the provided trip and traveler information. 
⚠️ IMPORTANT:
- Always include **all 12 mandatory sections** listed below in the exact order, even if minimal items are needed.
- Quantities, emojis, and items must be **dynamically decided** based on input data.
- Use **standardized emojis per section** consistently.
- Consider **traveler-specific details**: age, gender, medical notes, dietary restrictions, chronic conditions.
- Consider **trip details**: location, activities, duration, accommodation, budget, luggage style.
- Include optional, backup, and emergency items for all travelers.
- EXTREMELY IMPORTANT: Take the provided day-by-day temperatures VERY seriously. Evaluate the exact high and low bounds and precisely adjust clothing (thermal wear, summer wear, thin layers) to perfectly match the temperature swings. Do NOT recommend winter items for hot days, or summer items for freezing days.
- Ensure **practicality**: only include items travelers can realistically carry.
- Output **only the packing list in a flat JSON list**, do not include greetings, explanations, or nested arrays.


========================
MANDATORY SECTIONS (ALL CAPS)
========================
1. DOCUMENTS
2. CLOTHING
3. FOOTWEAR
4. TOILETRIES & PERSONAL CARE
5. ELECTRONICS & GADGETS
6. MEDICAL & HEALTH
7. ACCESSORIES
8. FOOD & SNACKS
9. ACTIVITY-SPECIFIC ITEMS
10. MISCELLANEOUS
11. WEATHER-SPECIFIC ITEMS
12. SAFETY & EMERGENCY ITEMS

========================
GENERAL INSTRUCTIONS
========================
- Include **quantities appropriate** for number of days, travelers, and laundry availability.
- Include **all traveler-specific adjustments**: age, gender, medical notes, dietary restrictions.
- Include **weather-appropriate items** based on destination temperature and season.
- Include **electronics and accessories** according to traveler details (phones, laptops, cameras, chargers).
- Include **food, snacks, and hydration** as per traveler type (solo, kids, elders, family).
- Include **backup and emergency items**: extra socks, ch+argers, first-aid, medications.
- For multi-person trips, include items for **each traveler individually**.
- **Do not use placeholders** like "stuff" or "if needed".
- Always use **consistent emojis per section**.
- Output must be a **flat JSON list**, not nested arrays.
- Ensure all items are **practical, ready-to-go, and realistic**.
- **Every section MUST have at least one item.** If no specific activities are provided, add general exploration items (e.g. "Daypack 🎒", "Comfortable walking wear 👕") to ACTIVITY-SPECIFIC ITEMS so it is never empty.

========================
OUTPUT FORMAT
========================
Respond ONLY with **JSON object** in this exact format:

{
  "packing_list": [
    "DOCUMENTS",
    "Passport 📄 (1 per traveler)",
    "Visa 🛂 (if required)",
    "Flight tickets 🎟️",
    "Hotel booking confirmation 🏨",
    "Travel insurance card 🏥",
    ...
    "CLOTHING",
    "T-shirts 👕 (N)",
    "Jeans / Pants 👖 (N)",
    "Shorts 🩳 (N)",
    "Sleepwear 😴 (N)",
    ...
    "FOOTWEAR",
    "Walking shoes 👟",
    "Sandals 🩴",
    ...
    "TOILETRIES & PERSONAL CARE",
    "Toothbrush 🪥",
    "Toothpaste 🪥",
    "Soap / Body wash 🧼",
    ...
    "MEDICAL & HEALTH",
    "First aid kit 🩹",
    "Prescription medication 💊",
    "Fever medication 💊",
    ...
    "ACCESSORIES",
    "Sunglasses 🕶️",
    "Backpack 🎒",
    "Umbrella ☔️",
    ...
    "ELECTRONICS & GADGETS",
    "Smartphone & charger 📱",
    "Laptop & charger 💻",
    "Camera 📷",
    ...
    "FOOD & SNACKS",
    "Non-perishable snacks 🍫",
    "Energy bars 🍿",
    "Reusable water bottle 💧",
    ...
    "ACTIVITY-SPECIFIC ITEMS",
    "Gym clothes 🏋️",
    "Swimwear 🏊",
    "Trekking shoes 🥾",
    ...
    "MISCELLANEOUS",
    "Notebook 📝",
    "Guidebook 📘",
    ...
    "WEATHER-SPECIFIC ITEMS",
    "Thermal wear 🧥 (cold weather)",
    "Sunscreen 🌞 (hot weather)",
    "Raincoat ☔️ (rainy weather)",
    ...
    "SAFETY & EMERGENCY ITEMS",
    "Travel locks 🔒",
    "Power bank 🔋",
    "Flashlight 🔦",
    "Emergency contact info 📇"
  ]
}
- Do not use nested arrays. Use a **flat list**.
- Each **section header is ALL CAPS, max 30 characters, no emojis**.
- Each **item under a section must include emoji** and quantity if relevant.
- Generate **all items dynamically**. Do not use placeholders like "if needed" or "stuff".
"""

    # Populate the dynamic data for the current request
    user_prompt = f"""
Location: {data['location']}
Duration: {data['days']} days
Trip Type: {data['trip_type']}
Purpose: {data['purpose']}
Activities: {data['activities']}
Stay Type: {data['stay_type']}
Budget: {data['budget']}
Food Preference: {data['food']}
Luggage Style: {data['luggage']}
Travel Mode: {data['travel_type']}
Travelers details: {data['people']}
Medical Notes: {data.get('medicalNotes')}
Dietary Notes: {data.get('dietaryNotes')}
Laundry Available: {data.get('laundry')}
Shopping Planned: {data.get('shopping')}
Photography Gear Needed: {data.get('photographyGear')}
Work Laptop: {data.get('workLaptop')}


Include all packing items dynamically based on:
- Traveler-specific medical notes, age, gender
- Weather conditions and temperature
- Trip duration, activities, and accommodation
- Budget and luggage style
- Ensure all 12 mandatory sections are included
- Use consistent emojis and realistic quantities

Environmental Context:
{temp_info}. Use this to determine weather-appropriate items.

Generate the JSON packing list now.
"""

    # Make the call to the Groq API model
    # We specify response_format={"type": "json_object"} to strictly enforce a JSON response structure.
    res = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt}
        ],
        temperature=0.4,
        max_completion_tokens=2000,
        response_format={"type": "json_object"}
    )
    
    try:
        content = res.choices[0].message.content
        
        try:
            result = json.loads(content)
        except json.JSONDecodeError as exc:
            logger.error(f"JSON decode error from Groq: {exc}\nRaw Content: {content[:200]}")
            raise HTTPException(status_code=500, detail="Invalid JSON response from AI")
        
        # Ensure packing_list items are clean from stray bullet points or numbering if the AI still added them
        clean_list = []
        for line in result.get("packing_list", []):
            if not isinstance(line, str):
                continue
            # Regex to remove bullets or numbering at the start of a response item
            clean_line = re.sub(r"^[•\-*\d.]+\s*", "", line).strip()
            if clean_line:
                clean_list.append(clean_line)
                
        logger.info(f"Successfully generated packing list containing {len(clean_list)} items")
        return {"packing_list": clean_list}
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Failed processing AI response: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Failed to generate packing list: {str(e)}")

# ==========================================
# ### API ENDPOINTS ###
# ==========================================

@app.post("/prefetch-weather", dependencies=[Depends(verify_api_key)])
def api_prefetch_weather(req: PrefetchWeatherRequest):
    """
    Endpoint to correct the city name using Groq, then prefetch the temperature.
    """
    client = Groq(api_key=GROQ_API_KEY)
    
    # 1. Correct city with Groq (Fastest model for simple NLP extraction)
    prompt = f"""
You are a strict location corrector.

Task:
Fix the spelling of the given city or country name.

Input:
{req.location}

Rules:
1. Return ONLY the corrected name.
2. Do NOT add any extra words, punctuation, or explanation.
3. Output must be a single valid city or country name.
4. If input is already correct, return it unchanged.
5. If the input is invalid or cannot be corrected, return exactly: INVALID.
6. Never return completely new city name as corected city name of user  
7. The corrected output MUST closely match the original input (only minor spelling fixes allowed).
8. Do NOT replace the input with a more popular or well-known city.
9. If unsure, return the original input exactly.
""" 
    try:
        res = client.chat.completions.create(
            model="llama-3.1-8b-instant",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.0,
            max_tokens=20
        )
        raw_response = res.choices[0].message.content.strip()
        
        # Fallback if Groq ignored instructions and wrote a conversational apology
        if len(raw_response.split()) > 4:
            corrected_city = req.location
        else:
            corrected_city = raw_response.strip(".,'\"")
            
    except Exception as e:
        logger.error(f"Groq city correction error: {e}")
        corrected_city = req.location # Fallback to original
        
    # 2. Fetch API forecast and populate cache for the corrected city
    temp = prefetch_and_cache_weather(corrected_city)
    
    return {"original": req.location, "location": corrected_city, "temperature": temp}


@app.post("/generate-packing-list", dependencies=[Depends(verify_api_key)])
# Node.js API gateway handles rate limiting
def api_generate_packing_list(request: Request, trip: TripRequestGenerate):
    """
    Primary API Endpoint to generate an AI-driven packing list.
    
    Expected Body (TripRequestGenerate format): 
      JSON with location, days, trip_type, budget, food, luggage, people, etc.
      
    Returns: 
      JSON response containing:
      {
         "packing_list": ["SECTION", "Item", "Item", ...]
      }
    """
    data = trip.dict()
    
    cache_key = json.dumps(data, sort_keys=True)
    current_time = time.time()
    
    # Check cache first to save AI calls
    if cache_key in generation_cache:
        cached_data, timestamp = generation_cache[cache_key]
        if current_time - timestamp < 300: # 5 minute TTL
            logger.info("Serving generated list from cache.")
            return cached_data
        else:
            del generation_cache[cache_key]

    ai_result = generate_packing_data(data)
    generation_cache[cache_key] = (ai_result, current_time)
    
    return ai_result

@app.post("/download-packing-list", dependencies=[Depends(verify_api_key)])
def api_download_packing_list(req: DownloadRequest):
    """
    Endpoint that packages the provided packing list directly into a downloadable .docx file.
    
    Expected Body: 
      JSON with packing_list.
      
    Returns:
      StreamingResponse serving a Word Document file as an attachment.
    """
    # Create DOCX and stream it back
    doc = create_docx(req.packing_list)
    return StreamingResponse(
        doc,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=Smart_Packing_List.docx"}
    )


class SuitcaseAnalysisRequest(BaseModel):
    """
    Expected input format for suitcase analyzer.
    """
    image_base64: str
    packing_list: list
    destination: str
    duration: int
    activities: str
    start_date: Optional[str] = None
    end_date: Optional[str] = None


class ChatRequest(BaseModel):
    message: str


# ==========================================
# ### VISION ANALYSIS HELPER FUNCTIONS ###
# ==========================================

def clean_item_name(item: str) -> str:
    """
    Strips emojis, ticks, and formatting symbols from an item's name.
    """
    import re
    # Remove surrogate pairs (non-BMP characters), other symbols, and variation selectors
    clean = re.sub(r'[^\u0000-\uFFFF]|[\u2600-\u27BF]|[\u2300-\u23FF]|[\u2B50-\u2B55]|[\u2934-\u2935]|[\uFE00-\uFE0F]', '', item)
    clean = re.sub(r'\s+', ' ', clean).strip()
    return clean


def build_vision_prompt(destination: str, weather: str, duration: int, activities: str, packing_list: list) -> str:
    """
    Builds the vision prompt instructions for the AI suitcase analyzer.
    """
    packing_list_str = "\n".join([f"- {item}" for item in packing_list])
    return f"""
You are an AI Travel suitcase capacity analyzer. You are given an image that may contain one or multiple suitcases/bags (up to 3 bags).
Analyze the suitcase(s) in the image and estimate:
1. The size of EACH suitcase/bag identified in the photo (e.g. Small, Medium, Large).
2. The approximate capacity of EACH suitcase/bag (e.g. in Liters).
3. The total combined capacity of all suitcases/bags shown in the photo.

Then, compare this total suitcase capacity against the traveler's details and proposed packing list:
- Destination: {destination}
- Weather/Temperature Forecast: {weather}
- Trip Duration: {duration} days
- Activities planned: {activities}

Here is the proposed packing list of items to categorize:
{packing_list_str}

Perform the following tasks:
1. Estimate the size and capacity of each bag, and state the combined total capacity (e.g. "Small (40L) and Medium (70L) - Total 110L"). Set "suitcase_size" to represent the overall sizing (e.g. "Medium", "Large", "2x Medium", "Small + Large" etc.) and "approximate_capacity" to represent the total combined capacity (e.g. "Total approx. 110 liters").
2. Assess whether the generated packing list can realistically fit in the identified suitcase(s) for the given duration and trip details. Provide a "comparison_summary".
3. Review each item from the proposed packing list and categorize it into one of these three lists:
   - "Must Pack": Items that are absolutely essential for the trip based on weather, duration, activities, and will fit.
   - "Optional": Items that are nice-to-have, but can be skipped if space is limited.
   - "Remove": Items that should be removed (e.g., because they are redundant, inappropriate for the weather, or won't fit in the suitcase).
4. For each categorized item, provide a "explanation" explaining why it is in that category.
5. EXTREMELY IMPORTANT: Categorize EVERY single item listed in the proposed packing list. Do not miss, omit, or rename any item. Strip all emojis and symbols from the item names in the final arrays (keep only name and quantity).
6. DISTRIBUTION RULE: You must realistically distribute the items across all three categories. Do not put all or almost all items in "Must Pack". Mark duplicate layers, extra toiletries, secondary footwear, and non-essential gadgets as "Optional". Mark items unsuitable for the forecast weather, duplicate primary items, or bulky items exceeding capacity constraints as "Remove".

You must respond ONLY with a JSON object of this structure:
{{
  "suitcase_size": "Description of suitcase size(s)",
  "approximate_capacity": "Total combined capacity description",
  "comparison_summary": "Your capacity review...",
  "categorized_items": {{
    "Must Pack": [
      {{"item": "cleaned item name (no emojis)", "explanation": "explanation string"}}
    ],
    "Optional": [
      {{"item": "cleaned item name (no emojis)", "explanation": "explanation string"}}
    ],
    "Remove": [
      {{"item": "cleaned item name (no emojis)", "explanation": "explanation string"}}
    ]
  }}
}}
Do not include any markup, markdown wrapper, or code fence. Just return the JSON object directly.
"""


def call_groq_vision_model(prompt: str, image_base64: str) -> dict:
    """
    Sends the prompt and base64-encoded image to the Groq Vision LLM.
    Attempts to query llama-4-maverick first, falling back to llama-4-scout if unavailable.
    """
    client = Groq(api_key=GROQ_API_KEY)
    
    # Ensure correct format for inline image data
    if not image_base64.startswith("data:image/"):
        image_url = f"data:image/jpeg;base64,{image_base64}"
    else:
        image_url = image_base64

    model_name = "meta-llama/llama-4-maverick-17b-128e-instruct"
    fallback_model = "meta-llama/llama-4-scout-17b-16e-instruct"

    try:
        logger.info(f"Attempting to query Groq vision model: {model_name}")
        res = client.chat.completions.create(
            model=model_name,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": prompt},
                        {
                            "type": "image_url",
                            "image_url": {
                                "url": image_url
                            }
                        }
                    ]
                }
            ],
            temperature=0.2,
            max_completion_tokens=2048,
            response_format={"type": "json_object"}
        )
        return res
    except Exception as e:
        if "model_not_found" in str(e) or "does not exist" in str(e) or "404" in str(e):
            logger.warning(f"Model '{model_name}' not available. Falling back to '{fallback_model}'. Error: {e}")
            res = client.chat.completions.create(
                model=fallback_model,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": image_url
                                }
                            }
                        ]
                    }
                ],
                temperature=0.2,
                max_completion_tokens=2048,
                response_format={"type": "json_object"}
            )
            return res
        else:
            logger.error(f"Error calling Groq vision model: {e}")
            raise


def parse_vision_response(raw_response: str) -> dict:
    """
    Parses and sanitizes the JSON response returned by the vision LLM.
    """
    try:
        cleaned = raw_response.strip()
        if cleaned.startswith("```json"):
            cleaned = cleaned[7:]
        elif cleaned.startswith("```"):
            cleaned = cleaned[3:]
        if cleaned.endswith("```"):
            cleaned = cleaned[:-3]
        cleaned = cleaned.strip()
        
        return json.loads(cleaned)
    except Exception as e:
        logger.error(f"JSON parsing error: {e}. Raw response: {raw_response}")
        raise HTTPException(status_code=500, detail="Invalid JSON response from vision model")


def optimize_packing_list(suitcase_info: dict, categorized_items: dict, original_items: list) -> dict:
    """
    Merges suitcase information and categorized packing list recommendations.
    Ensures 100% of original items are categorized without omission or duplicates,
    and strips any emojis/symbols from their names.
    """
    # Clean the category names
    must_pack_raw = categorized_items.get("Must Pack", [])
    optional_raw = categorized_items.get("Optional", [])
    remove_raw = categorized_items.get("Remove", [])
    
    # Help map items by matching their cleaned names (lowercase, stripped)
    def normalize(name):
        import re
        # Strip parenthesized text (like quantities (5) or (1 per traveler)) to prevent mismatch
        name_no_parens = re.sub(r'\(.*?\)', '', name)
        return "".join(c for c in name_no_parens.lower() if c.isalnum())
        
    # Build maps of LLM explanations
    explanations = {}
    llm_categories = {}
    
    for cat_name, items_list in [("Must Pack", must_pack_raw), ("Optional", optional_raw), ("Remove", remove_raw)]:
        for entry in items_list:
            if isinstance(entry, dict):
                item_name = entry.get("item", "")
                explanation = entry.get("explanation", "")
            else:
                item_name = str(entry)
                explanation = "Recommended recommendation based on suitcase capacity and trip details."
                
            clean_name = clean_item_name(item_name)
            norm_name = normalize(clean_name)
            if norm_name:
                explanations[norm_name] = explanation
                llm_categories[norm_name] = cat_name

    # Now, process every single original item
    final_categories = {
        "Must Pack": [],
        "Optional": [],
        "Remove": []
    }
    
    for item in original_items:
        clean_name = clean_item_name(item)
        norm_name = normalize(clean_name)
        
        # Retrieve explanation
        exp = explanations.get(norm_name, "Essential item for the trip details.")
        
        # Determine category (defaulting to Must Pack if LLM missed it)
        cat = llm_categories.get(norm_name, "Must Pack")
        
        final_categories[cat].append({
            "item": clean_name,
            "original_item": item, # Keep original string with emojis for frontend checkboxes
            "explanation": exp
        })
        
    return {
        "suitcase_size": suitcase_info.get("suitcase_size", "Medium"),
        "approximate_capacity": suitcase_info.get("approximate_capacity", "N/A"),
        "comparison_summary": suitcase_info.get("comparison_summary", ""),
        "categorized_items": final_categories
    }


# ==========================================
# ### VISION VALIDATION FUNCTIONS ###
# ==========================================

def build_validation_prompt() -> str:
    """
    Builds the vision validation prompt to verify the uploaded image.
    """
    return """
Analyze the uploaded image and verify whether it meets the following strict criteria:
1. The image contains exactly one (a single) suitcase or travel bag.
2. The suitcase or travel bag is empty and open.
3. The suitcase or travel bag is clearly visible and the photo is clear (not blurry, not obstructed, not closed).
4. The image DOES NOT contain people, animals, vehicles, random objects, unrelated scenes, multiple bags, or closed bags.

You must respond ONLY with a JSON object of this structure:
{
  "valid": true or false,
  "reason": "A short sentence explaining why the image is valid, or specifically why it is invalid."
}
Do not include any markup, markdown wrapper, or code fence. Just return the JSON object directly.
"""

def validate_suitcase_image(image_base64: str) -> dict:
    """
    Validates if the image contains a single, empty, open suitcase/travel bag.
    """
    prompt = build_validation_prompt()
    try:
        response = call_groq_vision_model(prompt, image_base64)
        content = response.choices[0].message.content
        parsed = parse_vision_response(content)
        logger.info(f"Vision validation parsed: {parsed}")
        return {
            "valid": bool(parsed.get("valid", False)),
            "reason": str(parsed.get("reason", "Unknown validation error."))
        }
    except Exception as e:
        logger.error(f"Error in validate_suitcase_image: {e}")
        return {
            "valid": False,
            "reason": "Failed to validate suitcase image due to model query error."
        }


def analyze_suitcase_image(image_base64: str, packing_list: list, destination: str, duration: int, activities: str, weather: str) -> dict:
    """
    Main controller for analyzing a suitcase image and matching its capacity.
    """
    prompt = build_vision_prompt(destination, weather, duration, activities, packing_list)
    res = call_groq_vision_model(prompt, image_base64)
    content = res.choices[0].message.content
    parsed = parse_vision_response(content)
    
    suitcase_info = {
        "suitcase_size": parsed.get("suitcase_size", "Medium"),
        "approximate_capacity": parsed.get("approximate_capacity", "N/A"),
        "comparison_summary": parsed.get("comparison_summary", "")
    }
    categorized = parsed.get("categorized_items", {
        "Must Pack": [],
        "Optional": [],
        "Remove": []
    })
    
    return optimize_packing_list(suitcase_info, categorized, packing_list)


@app.post("/analyze-suitcase", dependencies=[Depends(verify_api_key)])
def api_analyze_suitcase(req: SuitcaseAnalysisRequest):
    """
    Endpoint that processes the suitcase analyzer request.
    """
    # 1. Validate the uploaded image first
    validation = validate_suitcase_image(req.image_base64)
    if not validation.get("valid"):
        raise HTTPException(
            status_code=400,
            detail="Invalid image detected. Please upload a clear photo of your own empty, open suitcase or travel bag for analysis."
        )
        
    # 2. Fetch/compute weather
    weather_info = compute_full_trip_weather({
        "location": req.destination,
        "start_date": req.start_date,
        "end_date": req.end_date,
        "days": req.duration
    })
    
    # 3. Filter out category headers from packing list
    actual_items = [item for item in req.packing_list if not is_section_heading(item)]
    
    # 4. Run analysis
    result = analyze_suitcase_image(
        image_base64=req.image_base64,
        packing_list=actual_items,
        destination=req.destination,
        duration=req.duration,
        activities=req.activities,
        weather=weather_info
    )
    
    return result


# ==========================================
# ### RAG TRAVEL ASSISTANT CHATBOT FUNCTIONS ###
# ==========================================

# Path mappings for RAG
KNOWLEDGE_BASE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "knowledge_base")
CHROMA_DB_DIR = os.path.join(KNOWLEDGE_BASE_DIR, "chroma_db")


@app.post("/chat", dependencies=[Depends(verify_api_key)])
def api_chat(req: ChatRequest):
    """
    RAG-based Chatbot endpoint that accepts user questions,
    retrieves context from vector store, queries LLM, and returns the response.
    """
    try:
        from knowledge_base.rag import travel_chatbot
        return travel_chatbot(req.message)
    except Exception as e:
        logger.error(f"Error in api_chat: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/travel-chat", dependencies=[Depends(verify_api_key)])
def api_travel_chat(req: ChatRequest):
    """
    RAG-based Travel Chatbot endpoint that queries the vector database
    using custom embeddings and generates completions via Groq.
    """
    try:
        from knowledge_base.rag import travel_chatbot
        return travel_chatbot(req.message)
    except Exception as e:
        logger.error(f"Error in api_travel_chat: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ==========================================
# ### APPLICATION ENTRY POINT ###
# ==========================================
if __name__ == "__main__":
    import uvicorn
    # Start the server using Uvicorn. Enables hot-reload during active development.
    uvicorn.run(
        "main:app",
        host="0.0.0.0",
        port=int(os.environ.get("PORT", 5001)), # Changed to 5001 to prevent Node port collision
        reload=False
    )
